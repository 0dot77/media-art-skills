#!/usr/bin/env python3
"""
Markdown to DOCX converter for media-art-skills.

Usage:
    python md-to-docx.py input.md [output.docx]

If output path is omitted, saves as input.docx in the same directory.

Dependencies:
    pip install python-docx
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def parse_markdown_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Skip separator rows (|---|---|)
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    if i == 0:  # Header row
                        for run in paragraph.runs:
                            run.bold = True

    return table


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a .docx Word document."""
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    # Set heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "맑은 고딕"
        heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    i = 0
    table_buffer = []
    in_table = False
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]
            fmt = p.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            i += 1
            continue

        # Table handling
        if line.strip().startswith("|"):
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            rows = parse_markdown_table(table_buffer)
            if rows:
                add_table(doc, rows)
            table_buffer = []

        # Skip frontmatter
        if line.strip() == "---":
            if i == 0:
                # Skip YAML frontmatter
                i += 1
                while i < len(lines) and lines[i].strip() != "---":
                    i += 1
                i += 1
                continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        # Horizontal rule
        elif line.strip() == "---" or line.strip() == "***":
            doc.add_paragraph("─" * 50)
        # Bullet list
        elif line.strip().startswith("- "):
            text = line.strip()[2:]
            # Handle bold
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Bullet")
        # Numbered list
        elif re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s", "", line.strip())
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            p = doc.add_paragraph(text, style="List Number")
        # Blockquote
        elif line.strip().startswith("> "):
            text = line.strip()[2:]
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1)
            for run in p.runs:
                run.italic = True
        # Empty line
        elif line.strip() == "":
            pass
        # Regular paragraph
        else:
            text = line.strip()
            # Remove markdown bold/italic markers for clean display
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            if text:
                doc.add_paragraph(text)

        i += 1

    # Flush remaining table
    if table_buffer:
        rows = parse_markdown_table(table_buffer)
        if rows:
            add_table(doc, rows)

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md-to-docx.py input.md [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_suffix(".docx")

    result = md_to_docx(input_path, output_path)
    print(f"Saved: {result}")
